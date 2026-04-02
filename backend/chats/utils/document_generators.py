import csv
import io
from typing import Optional

from docx import Document
from docx.shared import Pt
from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
)
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from backend.chats.utils.text_utils import (
    normalize_text,
    extract_title,
    is_bullet_line,
    get_table_header_and_rows,
)

def generate_csv_from_text(text: str, user_title: Optional[str] = None):
    headers, rows = get_table_header_and_rows(text)
    title = extract_title(text, user_title)

    # utf-8-sig works better in Excel
    output = io.StringIO(newline="")
    writer = csv.writer(output)

    writer.writerow([title])
    writer.writerow([])
    writer.writerow(headers)

    for row in rows:
        writer.writerow(row)

    output.seek(0)
    return output

def generate_doc_from_text(text: str, user_title: Optional[str] = None):
    headers, rows = get_table_header_and_rows(text)
    title = extract_title(text, user_title)

    document = Document()

    heading = document.add_heading(title, 0)
    heading.alignment = 1

    subtitle = document.add_paragraph("Generated from user text")
    subtitle.alignment = 1

    if headers == ["Content"]:
        for row in rows:
            p = document.add_paragraph()
            run = p.add_run(row[0] if row else "")
            run.font.size = Pt(11)
    else:
        table = document.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for col, key in enumerate(headers):
            header_cells[col].text = str(key)

        for i, row in enumerate(rows, start=1):
            for col in range(len(headers)):
                value = row[col] if col < len(row) else ""
                table.rows[i].cells[col].text = str(value)

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _build_pdf_styles():
    styles = getSampleStyleSheet()

    if "CenterTitle" not in styles:
        styles.add(
            ParagraphStyle(
                name="CenterTitle",
                parent=styles["Title"],
                alignment=TA_CENTER,
                fontSize=18,
                leading=22,
                spaceAfter=12,
            )
     )

    if "SmallBody" not in styles:
        styles.add(
            ParagraphStyle(
                name="SmallBody",
                parent=styles["BodyText"],
                fontSize=9,
                leading=11,
                spaceAfter=4,
                alignment=TA_LEFT,
            )
        )

    if "SmallBullet" not in styles:
        styles.add(
            ParagraphStyle(
                name="SmallBullet",
                parent=styles["BodyText"],
                fontSize=10,
                leading=12,
                leftIndent=12,
                firstLineIndent=-8,
                spaceAfter=3,
            )
        )

    return styles

def generate_pdf_from_text(text: str, user_title: Optional[str] = None):
    headers, rows = get_table_header_and_rows(text)
    title = extract_title(text, user_title)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(letter),
        leftMargin=24,
        rightMargin=24,
        topMargin=24,
        bottomMargin=24,
    )

    styles = _build_pdf_styles()
    elements = []

    elements.append(Paragraph(title, styles["CenterTitle"]))
    elements.append(Paragraph("Generated from user text", styles["BodyText"]))
    elements.append(Spacer(1, 12))

    if headers == ["Content"]:
        for row in rows:
            text_line = row[0] if row else ""
            if is_bullet_line(text_line):
                text_line = re.sub(r"^\s*[-*•]\s+", "", text_line).strip()
            elements.append(Paragraph(text_line, styles["SmallBody"]))
            elements.append(Spacer(1, 4))

        doc.build(elements)
        buffer.seek(0)
        return buffer

    table_data = []
    table_data.append([Paragraph(str(h), styles["Heading5"]) for h in headers])

    for row in rows:
        row_cells = []
        for cell in row:
            row_cells.append(Paragraph(str(cell), styles["SmallBody"]))
        table_data.append(row_cells)

    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4B5563")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.75, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer
